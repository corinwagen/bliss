from __future__ import annotations

from dataclasses import dataclass, field
import bliss
import copy
import random
import docx
import os

random.seed(12345)


@dataclass
class SpeechRound:
    event: bliss.SpeechEvent
    name: str
    rooms: list[str]
    students: list[bliss.Student]
    judges_per_room: int = 1
    write_prefix: str = "write/"

    assignments: list[list[bliss.Student]] = field(init=False)

    def __post_init__(self):
        self.assignments = [[] for _ in self.rooms]

    @property
    def is_outround(self) -> bool:
        return self.name in ("QF", "SF", "F")

    def do_break(self, randomize=True, assignments=None):
        if self.broken():
            print("Already broken!")
            return

        if assignments:
            for i, room in enumerate(assignments):
                for student_id in room:
                    self.assignments[i].append(self.students[[s.id for s in self.students].index(student_id)])

        else:
            shuffled_students = copy.deepcopy(self.students)
            if randomize:
                random.shuffle(shuffled_students)
            else:
                print("not randomizing!")

            # assign students
            for idx, student in enumerate(shuffled_students):
                assignment_idx = idx % len(self.assignments)
                self.assignments[assignment_idx].append(student)

        print("Round broken!")

    def gen_ballots(self):
        for idx, room in enumerate(self.rooms):
            for judge_id in range(1, self.judges_per_room + 1):
                save_str = f"{self.write_prefix}{self.event.name}_Round{self.name}_Room{idx}_Ballot{judge_id}.docx"
                if os.path.exists(save_str):
                    print(f"Already file at {save_str} !")
                    return

                # from Walker
                temp = docx.Document(self.event.ballot_template)
                temp.paragraphs[2].text += self.name
                temp.paragraphs[4].text += f"{room} ({idx})"
                temp.paragraphs[6].text += f" ({judge_id})"
                table = temp.tables[0]

                for student_idx, student in enumerate(self.assignments[idx]):
                    table.cell(student_idx + 1, 0).text = f"{student.name} ({student.id})"

                temp.save(save_str)

    def gen_postings(self):
        save_str = f"{self.write_prefix}{self.event.name}_Round{self.name}_Postings.docx"
        if os.path.exists(save_str):
            print(f"Already file at {save_str} !")
            return

        # from Walker
        temp = docx.Document(self.event.posting_template)
        temp.paragraphs[3].text += self.name
        table = temp.tables[0]

        for idx, room in enumerate(self.rooms):
            table.cell(0, idx).text += f" {room} ({idx})"
            for student_idx, student in enumerate(self.assignments[idx]):
                table.cell(student_idx + 1, idx).text = f"{student.name} ({student.id})"

        temp.save(save_str)

    def broken(self):
        return any([len(a) for a in self.assignments])


#    def get_ranking_for_room(self, room_id):
#        # returns dict mapping from student ids to placing
#        ballots = self.ballots[room_id]
#
#        # presumes all ballots have same order
#        scores = np.zeros_like(ballots[0].student_ids, dtype=np.float64)
#        for ballot in ballots:
#            scores += 1/ballot.rankings
#
#        # awkward
#        scores_dict = {ballots[0].student_ids[i]: scores[i] for i in range(len(scores))}
#        sorted_scores = {k: v for k, v in sorted(scores_dict.items(), key=lambda item: item[1], reverse=True)}
#
#        # this is straight from stackoverflow
#        positions = {}
#        cur_score = None # Score we're examining
#        cur_count = 0 # Number of others that we've seen with this score
#
#        for ix, (name, score) in enumerate(sorted_scores.items()):
#            if score == cur_score: # Same score for this player as previous
#                cur_count += 1
#            else: # Different score from before
#                cur_score = score
#                cur_count = 0
#            positions[name] = ix - cur_count + 1 # Add 1 because ix is 0-based
#
#        return positions


@dataclass
class DebateRound:
    event: bliss.DebateEvent
    name: str
    rooms: list[str]
    teams: list[bliss.Team]
    judges_per_room: int = 1
    write_prefix: str = "write/"

    assignments: list[list[bliss.Student]] = field(init=False)
    ballots: list[bliss.DebateBallot] = field(init=False)

    def __post_init__(self):
        self.assignments = [[] for _ in self.rooms]
        self.ballots = [[] for _ in self.rooms]

    @property
    def is_outround(self) -> bool:
        return self.name in ("QF", "SF", "F")

    @property
    def students(self):
        studs = list()
        for team in self.teams:
            studs.append(team.students[0])
            studs.append(team.students[1])
        return studs

    def broken(self):
        return any([len(a) for a in self.assignments])

    def do_break(self, pairings=None, rankings=None):
        if self.broken():
            print("Already broken!")
            return

        if pairings is None:
            if rankings:
                # match 1 against 2, 3 against 4, and so forth
                shuffled_teams = [None] * len(self.teams)
                for idx, rank in enumerate(rankings):
                    shuffled_teams[rank] = self.teams[idx]
            else:
                shuffled_teams = copy.deepcopy(self.teams)
                random.shuffle(shuffled_teams)

            # assign students
            for idx, team in enumerate(shuffled_teams):
                assignment_idx = idx % len(self.assignments)
                self.assignments[assignment_idx].append(team)

        else:
            assert len(pairings) == len(self.rooms)
            for idx, room in enumerate(self.rooms):
                team1 = self.teams[pairings[idx][0]]
                team2 = self.teams[pairings[idx][1]]
                self.assignments[idx].append(team1)
                self.assignments[idx].append(team2)

        for idx, assignment in enumerate(self.assignments):
            # all this clever stuff is breaking somehow right now, and i don't have time to fix it.
            team1 = assignment[0]
            team2 = assignment[1]

            if team2 in team1.prev_seen:
                print(f"{team1.name} already met {team2.name}!")

            team1.prev_seen.append(team2)
            team2.prev_seen.append(team1)

            # make sure nobody gets all aff or all neg
            if team1.num_aff <= team2.num_aff:
                team1.num_aff += 1
            else:
                self.assignments[idx] = [team2, team1]
                team2.num_aff += 1

        #            print(f"{team1.name} ({team1.num_aff} aff) vs {team2.name} ({team2.num_aff} aff)")

        print("Round broken!")

    def gen_ballots(self):
        for idx, room in enumerate(self.rooms):
            for judge_id in range(1, self.judges_per_room + 1):
                save_str = f"{self.write_prefix}{self.event.name}_Round{self.name}_Room{idx}_Ballot{judge_id}.docx"
                if os.path.exists(save_str):
                    print(f"Already file at {save_str} !")
                    return

                temp = docx.Document(self.event.ballot_template)
                temp.paragraphs[2].text += self.name
                temp.paragraphs[4].text += f"{room} ({idx})"
                temp.paragraphs[6].text += f" ({judge_id})"

                team_table = temp.tables[0]
                student_table = temp.tables[1]
                student_counter = 1
                for team_idx, team in enumerate(self.assignments[idx]):
                    team_table.cell(team_idx + 1, 0).text = f"{team.name} ({team.id})"
                    for student in team.students:
                        student_table.cell(student_counter, 0).text = f"{student.name} ({student.id})"
                        student_counter += 1

                temp.save(save_str)

    def gen_postings(self):
        save_str = f"{self.write_prefix}{self.event.name}_Round{self.name}_Postings.docx"
        if os.path.exists(save_str):
            print(f"Already file at {save_str} !")
            return

        temp = docx.Document(self.event.posting_template)
        temp.paragraphs[3].text += self.name
        table = temp.tables[0]

        for idx, room in enumerate(self.rooms):
            table.cell(idx + 1, 0).text += f" {room} ({idx})"
            aff_team = self.assignments[idx][0]
            neg_team = self.assignments[idx][1]
            table.cell(idx + 1, 1).text = f"{aff_team.name} ({aff_team.id})"
            table.cell(idx + 1, 2).text = f"{neg_team.name} ({neg_team.id})"

        temp.save(save_str)
